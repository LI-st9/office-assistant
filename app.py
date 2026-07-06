# -*- coding: utf-8 -*-
"""
智能办公助手 - Gradio 
仿 DeepSeek 风格，极简居中单栏布局
支持：日常对话|pdf、word、excel等文档检索问答 | 生成Word/Excel | 修改样式| 流式输出
-智能意图分类和路由
-不同任务代理
-工具集成
-会话管理
-多快捷键
"""
__version__ = "1.0.0"
__author__ = "YL"

import os
import sys
import json
import json5
import re
import uuid
import time
import hashlib
import shutil
import argparse
from datetime import datetime
from typing import Iterator, Tuple, List, Dict, Any, Optional

import gradio as gr
import pandas as pd

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_community.document_loaders import PyMuPDFLoader, Docx2txtLoader, TextLoader
from qwen_agent.agents import Assistant
from qwen_agent.tools.base import BaseTool, register_tool
from qwen_agent.llm import get_chat_model

# ==========================================
# 配置
# ==========================================
class Config:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
    OUTPUT_DIR = os.path.join(BASE_DIR, "outputs")
    TEMP_DIR = os.path.join(BASE_DIR, "temp")
    MAX_FILE_SIZE = 50 * 1024 * 1024
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.txt', '.md'}
    SESSION_TIMEOUT = 3600

# 确保目录存在
os.makedirs(Config.UPLOAD_DIR, exist_ok=True)
os.makedirs(Config.OUTPUT_DIR, exist_ok=True)
os.makedirs(Config.TEMP_DIR, exist_ok=True)

# ==========================================
# 1. 配置与初始化
# ==========================================
os.environ["DASHSCOPE_API_KEY"] = os.getenv("DASHSCOPE_API_KEY", "your-api-key-here")

# PDF处理库
try:
    import fitz
    PDF_LOADER_AVAILABLE = True
except ImportError:
    PDF_LOADER_AVAILABLE = False

# OCR引擎
OCR_AVAILABLE = False
try:
    from paddleocr import PaddleOCR
    OCR_AVAILABLE = True
    ocr_engine = PaddleOCR(use_angle_cls=True, lang='ch')
except ImportError:
    OCR_AVAILABLE = False

# 全局分词器配置
TEXT_SPLITTER = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
CURRENT_FILE_PATH = None

# ==========================================
# 2. 公共转换函数
# ==========================================

WORD_FONT_SIZE_MAP = {
    '初号': 42, '小初': 36,
    '一号': 26, '小一': 24,
    '二号': 22, '小二': 18,
    '三号': 16, '小三': 15,
    '四号': 14, '小四': 12,
    '五号': 10.5, '小五': 9,
    '六号': 7.5, '小六': 6.5,
    '七号': 5.5, '八号': 5,
}

EXCEL_FONT_SIZE_MAP = {
    '初号': 36, '小初': 32,
    '一号': 26, '小一': 24,
    '二号': 22, '小二': 18,
    '三号': 16, '小三': 15,
    '四号': 14, '小四': 12,
    '五号': 10.5, '小五': 9,
    '六号': 7.5, '小六': 6.5,
}

def convert_word_font_size(size):
    if size is None:
        return 12
    if isinstance(size, (int, float)):
        return size
    if isinstance(size, str):
        num_match = re.match(r'^(\d+\.?\d*)', size)
        if num_match:
            return float(num_match.group(1))
        return WORD_FONT_SIZE_MAP.get(size, 12)
    return 12

def convert_excel_font_size(size):
    if size is None:
        return 11
    if isinstance(size, (int, float)):
        return max(6, min(48, size))
    if isinstance(size, str):
        num_match = re.match(r'^(\d+\.?\d*)', size)
        if num_match:
            return max(6, min(48, float(num_match.group(1))))
        return EXCEL_FONT_SIZE_MAP.get(size, 11)
    return 11

def clean_text_content(text):
    if not text:
        return text
    cleaned = re.sub(r'\n+', '\n', text)
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)
    return cleaned.strip()

def get_unique_filename(filename):
    if not os.path.exists(filename):
        return filename
    base, ext = os.path.splitext(filename)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{base}_{timestamp}{ext}"

def _parse_params(params):
    """增强容错的参数解析"""
    if isinstance(params, dict):
        return params
    if isinstance(params, str):
        try:
            return json5.loads(params)
        except Exception:
            try:
                start = params.find('{')
                end = params.rfind('}')
                if start != -1 and end != -1 and start < end:
                    json_str = params[start:end+1]
                    json_str = re.sub(r',\s*}', '}', json_str)
                    json_str = re.sub(r',\s*]', ']', json_str)
                    json_str = re.sub(r'//.*?$', '', json_str, flags=re.MULTILINE)
                    return json5.loads(json_str)
            except:
                pass
            return {}
    try:
        return json5.loads(json.dumps(params))
    except:
        return {}

# ==========================================
# 3. 文档处理引擎
# ==========================================
class DocumentProcessor:
    def __init__(self):
        self.embeddings = DashScopeEmbeddings(model="text-embedding-v3")

    def get_full_content(self, file_path):
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.pdf':
                return self._read_pdf_with_ocr(file_path)
            elif ext == '.docx':
                return self._read_docx(file_path)
            elif ext == '.xlsx':
                return self._read_excel(file_path)
            elif ext in ('.txt', '.md'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                return f"不支持的文件格式: {file_path}"
        except Exception as e:
            return f"读取文件失败: {str(e)}"

    def _read_pdf_with_ocr(self, file_path):
        extracted_text = ""
        if PDF_LOADER_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                full_text = ""
                for page_num, page in enumerate(doc):
                    page_text = page.get_text()
                    if page_text.strip():
                        full_text += page_text + "\n"
                    else:
                        blocks = page.get_text("blocks")
                        block_texts = []
                        for block in blocks:
                            if len(block) >= 5 and block[4].strip():
                                block_texts.append(block[4])
                        page_text = "\n".join(block_texts)
                        if page_text.strip():
                            full_text += page_text + "\n"
                        else:
                            html_text = page.get_text("html")
                            clean_text = re.sub(r'<[^>]+>', ' ', html_text)
                            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                            if clean_text:
                                full_text += clean_text + "\n"
                doc.close()
                if len(full_text.strip()) > 50:
                    print(f" PyMuPDF提取成功，共 {len(full_text.strip())} 字符")
                    return full_text
                else:
                    extracted_text = full_text
                    print(f"⚠️ PyMuPDF提取内容较少 ({len(full_text.strip())} 字符)")
            except Exception as e:
                print(f"PyMuPDF读取失败: {e}")

        try:
            import pdfplumber
            print("🔍 尝试使用 pdfplumber 提取...")
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        full_text += page_text + "\n"
                    else:
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                row_text = " | ".join([str(cell) for cell in row if cell])
                                if row_text:
                                    full_text += row_text + "\n"
                if len(full_text.strip()) > max(50, len(extracted_text.strip())):
                    print(f" pdfplumber提取成功，共 {len(full_text.strip())} 字符")
                    return full_text
                elif len(full_text.strip()) > len(extracted_text.strip()):
                    extracted_text = full_text
                    print(f"pdfplumber提取 {len(full_text.strip())} 字符")
        except ImportError:
            print("pdfplumber 未安装，跳过")
        except Exception as e:
            print(f"pdfplumber提取失败: {e}")

        if OCR_AVAILABLE:
            try:
                print("🔍 启动OCR识别...")
                result = ocr_engine.ocr(file_path)
                ocr_text = ""
                for page in result:
                    if page:
                        for line in page:
                            ocr_text += line[1][0] + "\n"
                if ocr_text.strip():
                    if len(ocr_text.strip()) > len(extracted_text.strip()) + 20:
                        print(f"✅ OCR识别成功，共 {len(ocr_text.strip())} 字符")
                        return ocr_text
                    else:
                        merged = extracted_text + "\n" + ocr_text
                        print(f"✅ 合并提取完成，共 {len(merged.strip())} 字符")
                        return merged
                else:
                    return "PDF内容无法识别，请检查文件是否损坏"
            except Exception as e:
                return f"OCR识别失败: {str(e)}"
        else:
            if extracted_text.strip():
                return extracted_text
            return "PDF解析失败：所有提取方法均无效"

    def _read_docx(self, file_path):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            content_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    header = "| " + " | ".join(table_data[0]) + " |"
                    separator = "| " + " | ".join(["---"] * len(table_data[0])) + " |"
                    rows = ["| " + " | ".join(row) + " |" for row in table_data[1:]]
                    content_parts.append("\n".join([header, separator] + rows))
            return "\n\n".join(content_parts)
        except Exception as e:
            return f"Word读取失败: {str(e)}"

    def _read_excel(self, file_path):
        try:
            df = pd.read_excel(file_path)
            return df.to_markdown(index=False)
        except Exception as e:
            return f"Excel读取失败: {str(e)}"

    def process_for_rag(self, file_path, query):
        try:
            if file_path.lower().endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
                if df.empty:
                    return "Excel文件为空"
                analysis_keywords = ['最大', '最小', '最高', '最低', '排名', '排序', '平均', '总和',
                                    '最后一名', '第一名', '第几名', '统计', '分析', '多少']
                if any(kw in query for kw in analysis_keywords):
                    return df.to_markdown(index=False)
                chunks = []
                for idx, row in df.iterrows():
                    content = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    chunks.append(f"【Excel数据-第{idx+1}行】\n{content}")
                docs = [Document(page_content=chunk) for chunk in chunks]
                vectorstore = FAISS.from_documents(docs, self.embeddings)
                results = vectorstore.similarity_search(query, k=6)
                return "\n\n---\n\n".join([doc.page_content for doc in results])

            loader = None
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.pdf':
                return self._read_pdf_with_ocr(file_path)
            elif ext == '.docx':
                loader = Docx2txtLoader(file_path)
            elif ext in ('.txt', '.md'):
                loader = TextLoader(file_path, encoding='utf-8')

            if not loader:
                return "不支持的文件格式"
            documents = loader.load()
            if not documents:
                return "文档内容为空"

            docs = TEXT_SPLITTER.split_documents(documents)
            if not docs:
                return "文档内容为空，无法检索"

            vectorstore = FAISS.from_documents(docs, self.embeddings)
            results = vectorstore.similarity_search(query, k=6)
            if not results:
                return "未找到与问题相关的内容"
            return "\n\n---\n\n".join([doc.page_content for doc in results])

        except Exception as e:
            return f"RAG检索失败: {str(e)}"

# ==========================================
# 4. 意图识别路由
# ==========================================
def route_intent(query: str) -> dict:
    classifier_prompt = """
    你是一个意图识别专家。请分析用户的指令，判断其任务类型。
    任务类型有三种：
    1. `format_task`: 包含提取信息生成新文件、格式转换、或者明确的'生成/创建'指令。
    2. `modify_task`: 包含'修改'、'调整'、'美化'等指令，且上下文包含文件路径。
    3. `qa_task`: 基于文档内容的问答、总结、分析等。

    请只返回一个 JSON 对象，包含 `task_type` 字段。
    """
    try:
        llm = get_chat_model({'model': 'qwen-turbo', 'model_server': 'dashscope'})
        messages = [{'role': 'system', 'content': classifier_prompt}, {'role': 'user', 'content': query}]
        response = []
        for chunk in llm.chat(messages=messages):
            response.extend(chunk)
        result_text = response[-1].get('content', '')

        try:
            task_info = json5.loads(result_text)
            task_type = task_info.get('task_type', 'qa_task')
        except:
            task_type = 'qa_task'
            if any(kw in result_text.lower() for kw in ['modify', '调整', '美化', '修改']):
                task_type = 'modify_task'
            elif any(kw in result_text.lower() for kw in ['format', '生成', '提取', '创建', '转成']):
                task_type = 'format_task'

        if task_type == 'qa_task':
            return {'task_type': task_type, 'model_config': {'model': 'qwen-turbo', 'model_server': 'dashscope'}}
        else:
            return {'task_type': task_type, 'model_config': {'model': 'qwen-max', 'model_server': 'dashscope'}}

    except Exception as e:
        print(f"意图识别失败: {e}，使用默认配置")
        return {'task_type': 'qa_task', 'model_config': {'model': 'qwen-turbo', 'model_server': 'dashscope'}}

# ==========================================
# 5. 工具注册
# ==========================================
doc_processor = DocumentProcessor()


@register_tool('process_excel')
class ProcessExcel(BaseTool):
    description = '对Excel表格数据进行排序、求和、求平均、计数等操作。支持数值排序和字符串排序。'
    parameters = [{
        'name': 'data', 'type': 'array', 'description': '表格数据（列表字典）', 'required': True
    }, {
        'name': 'operation', 'type': 'string', 'description': '操作类型: sort/sum/average/count', 'required': True
    }, {
        'name': 'column', 'type': 'string', 'description': '操作的目标列名', 'required': True
    }, {
        'name': 'ascending', 'type': 'boolean', 'description': '排序时是否升序（默认True）', 'required': False
    }, {
        'name': 'new_column_name', 'type': 'string', 'description': 'sum/average/count时，结果列的名称', 'required': False
    }]

    def call(self, params, **kwargs) -> str:
        params = _parse_params(params)
        try:
            data = params['data']
            operation = params['operation']
            column = params['column']

            df = pd.DataFrame(data)

            if column not in df.columns:
                return json.dumps({"status": "error", "message": f"列 '{column}' 不存在"})

            result = {}

            if operation == 'sort':
                ascending = params.get('ascending', True)

                try:
                    numeric_series = pd.to_numeric(df[column], errors='coerce')
                    valid_ratio = numeric_series.notna().sum() / len(numeric_series) if len(numeric_series) > 0 else 0
                    if valid_ratio > 0.8:
                        df['_sort_key'] = numeric_series
                        sorted_df = df.sort_values(by='_sort_key', ascending=ascending)
                        sorted_df = sorted_df.drop(columns=['_sort_key'])
                    else:
                        sorted_df = df.sort_values(by=column, ascending=ascending)
                except:
                    sorted_df = df.sort_values(by=column, ascending=ascending)

                result = {
                    'operation': 'sort',
                    'column': column,
                    'ascending': ascending,
                    'data': sorted_df.to_dict('records')
                }

            elif operation == 'sum':
                numeric_col = pd.to_numeric(df[column], errors='coerce')
                if numeric_col.isna().all():
                    return json.dumps({"status": "error", "message": f"列 '{column}' 无法转换为数值进行求和"})
                total = numeric_col.sum()
                result = {
                    'operation': 'sum',
                    'column': column,
                    'result': total,
                    'data': data
                }

            elif operation == 'average':
                numeric_col = pd.to_numeric(df[column], errors='coerce')
                if numeric_col.isna().all():
                    return json.dumps({"status": "error", "message": f"列 '{column}' 无法转换为数值进行求平均"})
                avg = numeric_col.mean()
                result = {
                    'operation': 'average',
                    'column': column,
                    'result': avg,
                    'data': data
                }

            elif operation == 'count':
                count = df[column].count()
                result = {
                    'operation': 'count',
                    'column': column,
                    'result': count,
                    'data': data
                }

            else:
                return json.dumps({"status": "error", "message": f"不支持的操作: {operation}"})

            return json.dumps({"status": "success", "result": result})

        except Exception as e:
            return json.dumps({"status": "error", "message": str(e)})


@register_tool('modify_word')
class ModifyWord(BaseTool):
    description = '修改原Word文档的样式。可指定修改标题样式或正文样式。'
    parameters = [{
        'name': 'file_path', 'type': 'string', 'description': '原Word文件路径', 'required': True
    }, {
        'name': 'save_as', 'type': 'string', 'description': '新文件名（可选，不指定则覆盖原文件）', 'required': False
    }, {
        'name': 'target_style', 'type': 'string', 'description': '要修改的样式: "title"(标题) / "body"(正文) / "all"(全部)', 'required': True
    }, {
        'name': 'font_name', 'type': 'string', 'description': '新字体', 'required': False
    }, {
        'name': 'font_size', 'type': 'string', 'description': '新字号（支持中文字号或数字）', 'required': False
    }, {
        'name': 'bold', 'type': 'boolean', 'description': '是否加粗', 'required': False
    }, {
        'name': 'italic', 'type': 'boolean', 'description': '是否斜体', 'required': False
    }, {
        'name': 'underline', 'type': 'boolean', 'description': '是否下划线', 'required': False
    }, {
        'name': 'color', 'type': 'string', 'description': '颜色（如 #FF0000 或 red）', 'required': False
    }, {
        'name': 'align', 'type': 'string', 'description': '对齐方式: left/center/right/justify', 'required': False
    }, {
        'name': 'indent', 'type': 'number', 'description': '首行缩进（字符数）', 'required': False
    }, {
        'name': 'line_spacing', 'type': 'number', 'description': '行间距（倍数，如1.5）', 'required': False
    }]

    def call(self, params, **kwargs) -> str:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        params = _parse_params(params)
        try:
            file_path = params['file_path']
            save_as = params.get('save_as')
            target_style = params.get('target_style', 'all')

            if save_as:
                if not save_as.endswith('.docx'):
                    save_as += '.docx'
                base_dir = os.path.dirname(file_path)
                new_path = os.path.join(base_dir, save_as)
                new_path = get_unique_filename(new_path)
                shutil.copy2(file_path, new_path)
                work_path = new_path
                backup_path = None
            else:
                backup_path = file_path + '.bak'
                shutil.copy2(file_path, backup_path)
                work_path = file_path

            doc = DocxDocument(work_path)

            font_name = params.get('font_name')
            font_size_raw = params.get('font_size')
            font_size = convert_word_font_size(font_size_raw) if font_size_raw is not None else None
            bold = params.get('bold')
            italic = params.get('italic')
            underline = params.get('underline')
            color = params.get('color')
            align = params.get('align')
            indent = params.get('indent')
            line_spacing = params.get('line_spacing')

            align_map = {
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }

            def parse_color(color_str):
                if not color_str:
                    return None
                if color_str.startswith('#'):
                    try:
                        r, g, b = tuple(int(color_str[i:i+2], 16) for i in (1, 3, 5))
                        return RGBColor(r, g, b)
                    except:
                        return None
                color_map = {
                    'red': RGBColor(255, 0, 0),
                    'blue': RGBColor(0, 0, 255),
                    'green': RGBColor(0, 128, 0),
                    'black': RGBColor(0, 0, 0),
                    'white': RGBColor(255, 255, 255)
                }
                return color_map.get(color_str)

            def set_indent(paragraph, indent_chars, font_size_override=None):
                if indent_chars is None:
                    return
                para_font_size = font_size_override
                if para_font_size is None:
                    for run in paragraph.runs:
                        if run.font.size is not None:
                            para_font_size = run.font.size.pt
                            break
                if para_font_size is None:
                    para_font_size = 12
                indent_pt = indent_chars * para_font_size
                paragraph.paragraph_format.first_line_indent = Pt(indent_pt)
                try:
                    p = paragraph._element
                    pPr = p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p.insert(0, pPr)
                    old_ind = pPr.find(qn('w:ind'))
                    if old_ind is not None:
                        pPr.remove(old_ind)
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:firstLine'), str(int(indent_pt * 20)))
                    pPr.append(ind)
                except Exception as e:
                    print(f"⚠️ XML缩进设置失败: {e}")

            para_count = 0
            for para in doc.paragraphs:
                is_title = para.style.name.startswith('Heading') or para.style.name == 'Title'
                if target_style == 'title' and not is_title:
                    continue
                if target_style == 'body' and is_title:
                    continue
                para_count += 1
                if indent is not None:
                    set_indent(para, indent, font_size)
                if line_spacing is not None:
                    para.paragraph_format.line_spacing = line_spacing
                if align is not None and align in align_map:
                    para.alignment = align_map[align]
                for run in para.runs:
                    if font_name:
                        run.font.name = font_name
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    if font_size is not None:
                        run.font.size = Pt(font_size)
                    if bold is not None:
                        run.bold = bold
                    if italic is not None:
                        run.italic = italic
                    if underline is not None:
                        run.underline = underline
                    if color:
                        rgb = parse_color(color)
                        if rgb:
                            run.font.color.rgb = rgb

            print(f"✅ 已处理 {para_count} 个段落")
            doc.save(work_path)

            if save_as:
                msg = f"✅ 修改完成，共处理 {para_count} 个段落，新文件已保存为: {work_path}"
            else:
                msg = f"✅ 修改完成，共处理 {para_count} 个段落，原文件已备份为: {backup_path}"

            return json.dumps({"status": "success", "message": msg, "file": work_path})
        except Exception as e:
            return json.dumps({"status": "error", "message": str(e)})


@register_tool('create_word')
class CreateWord(BaseTool):
    description = '根据用户需求生成一个新的Word文档，支持丰富样式（支持中文字号、首行缩进、行间距等）。'
    parameters = [{
        'name': 'content', 'type': 'array', 'description': '内容列表', 'required': True
    }, {
        'name': 'filename', 'type': 'string', 'description': '输出文件名', 'required': True
    }, {
        'name': 'default_font', 'type': 'string', 'description': '默认字体', 'required': False
    }, {
        'name': 'default_size', 'type': 'string', 'description': '默认字号（支持中文字号如五号，或数字如12）', 'required': False
    }, {
        'name': 'default_line_spacing', 'type': 'number', 'description': '默认行间距（如1.5）', 'required': False
    }]

    def call(self, params, **kwargs) -> str:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        params = _parse_params(params)

        if 'content' in params:
            cleaned_content = []
            for item in params['content']:
                if isinstance(item, str):
                    cleaned = clean_text_content(item)
                    if cleaned:
                        cleaned_content.append(cleaned)
                elif isinstance(item, dict) and 'text' in item:
                    cleaned = clean_text_content(item['text'])
                    if cleaned:
                        new_item = item.copy()
                        new_item['text'] = cleaned
                        cleaned_content.append(new_item)
                else:
                    cleaned_content.append(item)
            params['content'] = cleaned_content

        try:
            doc = DocxDocument()

            default_font = params.get('default_font', '宋体')
            default_size_raw = params.get('default_size', 12)
            default_size = convert_word_font_size(default_size_raw)
            default_line_spacing = params.get('default_line_spacing', 1.5)

            for item in params['content']:
                if isinstance(item, str):
                    p = doc.add_paragraph(item)
                    for run in p.runs:
                        run.font.name = default_font
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), default_font)
                        run.font.size = Pt(default_size)
                    p.paragraph_format.line_spacing = default_line_spacing
                    continue

                text = item.get('text', '')
                style = item.get('style', 'normal')

                if style in ['heading1', 'heading2', 'heading3']:
                    level = int(style[-1])
                    heading = doc.add_heading(level=level)
                    run = heading.add_run(text)
                    if item.get('font'):
                        run.font.name = item['font']
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), item['font'])
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(text)

                    font_name = item.get('font', default_font)
                    run.font.name = font_name
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

                    font_size_raw = item.get('font_size', default_size)
                    font_size = convert_word_font_size(font_size_raw)
                    run.font.size = Pt(font_size)

                    run.bold = item.get('bold', False)
                    run.italic = item.get('italic', False)
                    run.underline = item.get('underline', False)

                    color = item.get('color')
                    if color:
                        try:
                            if color.startswith('#'):
                                r, g, b = tuple(int(color[i:i+2], 16) for i in (1, 3, 5))
                                run.font.color.rgb = RGBColor(r, g, b)
                            elif color in ['red', 'blue', 'green', 'black', 'white']:
                                color_map = {'red': RGBColor(255,0,0), 'blue': RGBColor(0,0,255),
                                           'green': RGBColor(0,128,0), 'black': RGBColor(0,0,0)}
                                if color in color_map:
                                    run.font.color.rgb = color_map[color]
                        except:
                            pass

                    align = item.get('align', 'left')
                    align_map = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT,
                                'left': WD_ALIGN_PARAGRAPH.LEFT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
                    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

                    indent = item.get('indent')
                    if indent is not None:
                        p.paragraph_format.first_line_indent = Pt(font_size * indent)

                    line_spacing = item.get('line_spacing', default_line_spacing)
                    p.paragraph_format.line_spacing = line_spacing

            filename = params['filename']
            if not filename.endswith('.docx'):
                filename += '.docx'
            
            # 保存到输出目录
            filename = os.path.join(Config.OUTPUT_DIR, filename)
            filename = get_unique_filename(filename)

            doc.save(filename)
            return json.dumps({"status": "success", "file": filename, "message": f"✅ Word文件已生成: {filename}"})
        except Exception as e:
            return json.dumps({"status": "error", "message": str(e)})


@register_tool('create_excel')
class CreateExcel(BaseTool):
    description = '根据用户需求生成一个新的Excel文件，支持丰富样式（支持中文字号、水平/垂直对齐、行高、列宽等）。'
    parameters = [{
        'name': 'data', 'type': 'array', 'description': '表格数据', 'required': True
    }, {
        'name': 'filename', 'type': 'string', 'description': '输出文件名', 'required': True
    }, {
        'name': 'font_name', 'type': 'string', 'description': '字体（如 宋体、Arial）', 'required': False
    }, {
        'name': 'font_size', 'type': 'string', 'description': '字号（支持数字如12，或中文如五号）', 'required': False
    }, {
        'name': 'header_color', 'type': 'string', 'description': '表头背景色（如 #4F81BD）', 'required': False
    }, {
        'name': 'header_text_color', 'type': 'string', 'description': '表头文字颜色（如 #FFFFFF）', 'required': False
    }, {
        'name': 'row_height', 'type': 'number', 'description': '行高（单位：磅，如 20）', 'required': False
    }, {
        'name': 'column_widths', 'type': 'object', 'description': '列宽（单位：字符数，如 {"姓名": 10}）', 'required': False
    }, {
        'name': 'horizontal_align', 'type': 'string', 'description': '水平对齐: left/center/right', 'required': False
    }, {
        'name': 'vertical_align', 'type': 'string', 'description': '垂直对齐: top/center/bottom', 'required': False
    }]

    def call(self, params, **kwargs) -> str:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter

        params = _parse_params(params)
        try:
            df = pd.DataFrame(params['data'])
            filename = params['filename']
            if not filename.endswith('.xlsx'):
                filename += '.xlsx'

            #  保存到输出目录
            filename = os.path.join(Config.OUTPUT_DIR, filename)
            filename = get_unique_filename(filename)

            df.to_excel(filename, index=False)

            wb = openpyxl.load_workbook(filename)
            ws = wb.active

            font_name = params.get('font_name', '宋体')
            font_size_raw = params.get('font_size', 11)
            font_size = convert_excel_font_size(font_size_raw)

            header_color = params.get('header_color', '4F81BD')
            header_text_color = params.get('header_text_color', 'FFFFFF')

            horizontal_align = params.get('horizontal_align', 'center')
            vertical_align = params.get('vertical_align', 'center')
            h_align_map = {'left': 'left', 'center': 'center', 'right': 'right'}
            v_align_map = {'top': 'top', 'center': 'center', 'bottom': 'bottom'}
            h_align = h_align_map.get(horizontal_align, 'center')
            v_align = v_align_map.get(vertical_align, 'center')

            row_height = params.get('row_height')

            border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            for row_idx, row in enumerate(ws.iter_rows(), 1):
                for cell in row:
                    cell.font = Font(name=font_name, size=font_size)
                    cell.alignment = Alignment(horizontal=h_align, vertical=v_align)
                    cell.border = border
                if row_height:
                    ws.row_dimensions[row_idx].height = row_height

            header_font = Font(name=font_name, size=font_size, bold=True, color=header_text_color)
            header_fill = PatternFill(start_color=header_color, end_color=header_color, fill_type="solid")

            for cell in ws[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal=h_align, vertical=v_align)
                cell.border = border

            if params.get('column_widths'):
                for col_name, width in params['column_widths'].items():
                    if col_name in df.columns:
                        col_idx = df.columns.get_loc(col_name) + 1
                        ws.column_dimensions[get_column_letter(col_idx)].width = width
            else:
                for col in ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 30)
                    ws.column_dimensions[column].width = adjusted_width

            wb.save(filename)
            return json.dumps({"status": "success", "file": filename, "message": f"✅ Excel文件已生成: {filename}"})
        except Exception as e:
            return json.dumps({"status": "error", "message": str(e)})


@register_tool('modify_excel')
class ModifyExcel(BaseTool):
    description = '在原Excel文件基础上修改样式（支持中文字号、水平/垂直对齐、行高、列宽等）并保存。'
    parameters = [{
        'name': 'file_path', 'type': 'string', 'description': '原文件路径', 'required': True
    }, {
        'name': 'font_name', 'type': 'string', 'description': '字体', 'required': False
    }, {
        'name': 'font_size', 'type': 'string', 'description': '字号（支持数字如12，或中文如五号）', 'required': False
    }, {
        'name': 'header_color', 'type': 'string', 'description': '表头背景色', 'required': False
    }, {
        'name': 'row_height', 'type': 'number', 'description': '行高（单位：磅）', 'required': False
    }, {
        'name': 'column_widths', 'type': 'object', 'description': '列宽设置（单位：字符数）', 'required': False
    }, {
        'name': 'horizontal_align', 'type': 'string', 'description': '水平对齐: left/center/right', 'required': False
    }, {
        'name': 'vertical_align', 'type': 'string', 'description': '垂直对齐: top/center/bottom', 'required': False
    }]

    def call(self, params, **kwargs) -> str:
        import openpyxl
        from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
        from openpyxl.utils import get_column_letter

        params = _parse_params(params)
        try:
            file_path = params['file_path']
            backup_path = file_path + '.bak'
            shutil.copy2(file_path, backup_path)

            wb = openpyxl.load_workbook(file_path)
            ws = wb.active

            font_name = params.get('font_name')
            font_size_raw = params.get('font_size')
            font_size = convert_excel_font_size(font_size_raw) if font_size_raw is not None else None
            header_color = params.get('header_color')
            row_height = params.get('row_height')

            horizontal_align = params.get('horizontal_align', 'center')
            vertical_align = params.get('vertical_align', 'center')
            h_align_map = {'left': 'left', 'center': 'center', 'right': 'right'}
            v_align_map = {'top': 'top', 'center': 'center', 'bottom': 'bottom'}
            h_align = h_align_map.get(horizontal_align, 'center')
            v_align = v_align_map.get(vertical_align, 'center')

            border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            for row_idx, row in enumerate(ws.iter_rows(), 1):
                for cell in row:
                    if font_name or font_size is not None:
                        current_font = Font(
                            name=font_name or '宋体',
                            size=font_size if font_size is not None else 11
                        )
                        cell.font = current_font
                    cell.alignment = Alignment(horizontal=h_align, vertical=v_align)
                    cell.border = border
                if row_height:
                    ws.row_dimensions[row_idx].height = row_height

            if header_color:
                header_fill = PatternFill(start_color=header_color, end_color=header_color, fill_type="solid")
                for cell in ws[1]:
                    cell.fill = header_fill

            if params.get('column_widths'):
                for col_name, width in params['column_widths'].items():
                    for cell in ws[1]:
                        if cell.value == col_name:
                            ws.column_dimensions[get_column_letter(cell.column)].width = width
                            break

            wb.save(file_path)
            return json.dumps({"status": "success", "message": f"Excel修改完成，原文件已备份为 {backup_path}"})
        except Exception as e:
            return json.dumps({"status": "error", "message": str(e)})


# ==========================================
# 6. Agent 创建函数
# ==========================================
def create_agent_for_task(task_type, system_prompt):
    if task_type == 'qa_task':
        llm_cfg = {'model': 'qwen-turbo', 'model_server': 'dashscope'}
        return Assistant(llm=llm_cfg, system_message=system_prompt)
    else:
        llm_cfg = {'model': 'qwen-max', 'model_server': 'dashscope'}
        return Assistant(
            llm=llm_cfg,
            system_message=system_prompt,
            function_list=['process_excel', 'create_word', 'create_excel', 'modify_word', 'modify_excel']
        )


# ==========================================
# 7. 会话管理
# ==========================================
class SessionManager:
    def __init__(self, timeout_seconds: int = 3600):
        self.sessions: Dict[str, Dict] = {}
        self.timeout = timeout_seconds
        self._cleanup_interval = 300
        self._last_cleanup = time.time()

    def get_or_create_session(self, session_id: Optional[str] = None) -> Tuple[str, Dict]:
        if session_id and session_id in self.sessions:
            session = self.sessions[session_id]
            session['last_active'] = time.time()
            return session_id, session

        # 如果传入的 session_id 不存在，使用传入的 id 创建新 session
        if session_id:
            new_session_id = session_id
        else:
            new_session_id = str(uuid.uuid4())[:8]
            
        session = {
            'id': new_session_id,
            'messages': [],
            'current_file_path': None,
            'current_file_info': None,
            'doc_processor': DocumentProcessor(),
            'created_at': time.time(),
            'last_active': time.time(),
        }
        self.sessions[new_session_id] = session
        self._cleanup()
        return new_session_id, session

    def update_session(self, session_id: str, updates: Dict):
        if session_id in self.sessions:
            self.sessions[session_id].update(updates)
            self.sessions[session_id]['last_active'] = time.time()

    def get_session(self, session_id: str) -> Optional[Dict]:
        if session_id in self.sessions:
            self.sessions[session_id]['last_active'] = time.time()
            return self.sessions[session_id]
        return None

    def get_chat_history(self, session_id: str) -> List[Dict]:
        """获取完整的对话历史"""
        session = self.get_session(session_id)
        if session:
            return session.get('messages', [])
        return []

    def _cleanup(self):
        now = time.time()
        if now - self._last_cleanup < self._cleanup_interval:
            return
        expired = [sid for sid, session in self.sessions.items()
                   if now - session['last_active'] > self.timeout]
        for sid in expired:
            del self.sessions[sid]
        self._last_cleanup = now


# ==========================================
# 8. 核心业务逻辑封装
# ==========================================
class OfficeAssistant:
    def __init__(self):
        self.session_manager = SessionManager()

    def upload_file(self, file, session_id: str) -> Tuple[str, str, str]:
        if file is None:
            return "未上传文件", "📎 未加载", "0 字"

        _, session = self.session_manager.get_or_create_session(session_id)
        print(f"📤 上传文件 - Session ID: {session_id}")
        print(f"📤 上传文件 - Session 对象 ID: {id(session)}")

        try:
            session_dir = os.path.join(Config.UPLOAD_DIR, session_id)
            os.makedirs(session_dir, exist_ok=True)

            if hasattr(file, 'name'):
                original_name = os.path.basename(file.name)
            else:
                original_name = os.path.basename(file)

            saved_path = os.path.join(session_dir, original_name)
            if hasattr(file, 'name'):
                shutil.copy2(file.name, saved_path)
            else:
                shutil.copy2(file, saved_path)

            doc_processor = session['doc_processor']
            content = doc_processor.get_full_content(saved_path)

            print(f"📁 文件已保存到 session: {saved_path}")
            print(f"📁 session['current_file_path']: {session.get('current_file_path')}")

            if "失败" in content or "不支持" in content:
                return f"⚠️ {content}", "📎 加载失败", "0 字"

            word_count = len(content)
            file_size = os.path.getsize(saved_path)
            size_str = f"{file_size / 1024:.1f} KB" if file_size < 1024 * 1024 else f"{file_size / (1024 * 1024):.1f} MB"

            session['current_file_path'] = saved_path
            session['current_file_info'] = {
                'name': original_name,
                'size': size_str,
                'word_count': word_count
            }

            self.session_manager.update_session(session_id, {
                'current_file_path': saved_path,
                'current_file_info': session['current_file_info']
            })
            verify_session = self.session_manager.get_session(session_id)
            print(f"📁 验证 session['current_file_path']: {verify_session.get('current_file_path') if verify_session else 'None'}")

            status = f"📎 {original_name}  |  {size_str}  |  {word_count} 字"
            return f"✅ 加载成功: {original_name}", status, f"{word_count} 字"

        except Exception as e:
            return f"⚠️ 上传失败: {str(e)}", "📎 加载失败", "0 字"

    def process_message(
        self,
        message: str,
        chat_history: List[Dict[str, str]], 
        session_id: str
    ) -> Iterator[Tuple[str, str, str, str, str]]:
        if not message or not message.strip():
            yield "请输入有效指令", "待识别", "qwen-turbo", "0", ""
            return

        _, session = self.session_manager.get_or_create_session(session_id)
        file_path = session.get('current_file_path')

        # 打印调试信息

        
        print(f"📁 process_message - 当前文件路径: {file_path}")
        print(f"📁 session ID: {session_id}")
        print(f"📁 session 对象 ID: {id(session)}")
        print(f"💬 chat_history 条数: {len(chat_history)}")

        task_type = "qa_task"
        model_name = "qwen-turbo"
        generated_file = ""

        try:
            yield f"🔍 分析任务类型...", "识别中", "qwen-turbo", f"{len(session['messages']) + 1}", ""

            intent_result = route_intent(message)
            task_type = intent_result['task_type']
            model_name = intent_result['model_config']['model']

            task_label = {
                'qa_task': '问答/总结',
                'format_task': '生成/创建',
                'modify_task': '修改/美化'
            }.get(task_type, '未知')

            yield f"📌 识别为: {task_label}，使用 {model_name}", task_label, model_name, f"{len(session['messages']) + 1}", ""

            context = ""
            if file_path and os.path.exists(file_path):
                if task_type == 'qa_task':
                    context = session['doc_processor'].process_for_rag(file_path, message)
                else:
                    context = session['doc_processor'].get_full_content(file_path)

                if context and ("失败" in context or "不支持" in context):
                    yield f"⚠️ {context}", task_label, model_name, f"{len(session['messages']) + 1}", ""
                    return
            if task_type == 'qa_task':
                system_prompt = f"""
    你是一个专业的智能办公助手。

    【文档内容】:
    {context if context else "（无文档加载，请进行日常对话）"}

    【任务要求】:
    - 请基于以上【文档内容】回答用户的问题
    - 如果是数据表格，请准确分析数据并给出答案
    - 如果文档内容中没有相关信息，请如实告知，禁止编造
    - 回答要简洁、准确、有条理
    """
            else:
                system_prompt = f"""
    你是一个专业的智能办公助手。

    【文档内容】:
    {context if context else "（无文档加载，请进行日常对话）"}

    【核心工作流程 - 从文档提取数据生成表格】:

    当用户要求从文档中提取数据并生成Excel时，请按以下步骤操作：

    1. 阅读【文档内容】，理解文档结构和用户需求
    2. 根据用户需求，识别需要提取的字段（如：姓名、成绩、学号等）
    3. 从文档中逐条提取数据，构建 JSON 数组
    4. 调用 create_excel 工具生成表格

    【数据提取格式要求】:
    提取的数据必须是 JSON 数组，格式为 [{{"列名1": 值1, "列名2": 值2}}, ...]
    不要添加任何额外文字
    如果数据量大，可以分批生成
    调用 create_excel 时，data 参数必须是 JSON 数组

    【示例 - 提取数据生成Excel】:
    用户: "从这份文档中提取学生姓名和成绩，生成Excel"
    文档内容: "张三 90分，李四 85分，王五 95分"

    你的工作:
    1. 识别字段: ["姓名", "成绩"]
    2. 提取数据: [{{"姓名": "张三", "成绩": 90}}, {{"姓名": "李四", "成绩": 85}}, {{"姓名": "王五", "成绩": 95}}]
    3. 调用 create_excel({{"data": [...], "filename": "学生成绩.xlsx"}})

    【Word样式设置 - 重要示例】:

    示例1 - 设置标题为黑体三号居中：
    {{
      "content": [
        {{"text": "这是标题", "font": "黑体", "font_size": "三号", "align": "center"}},
        {{"text": "这是正文内容", "font": "宋体", "font_size": "五号"}}
      ],
      "filename": "文档.docx"
    }}

    ⚠️ 重要：首行缩进必须在每个段落单独设置！
    示例2 - 正文段落首行缩进2字符（每个段落都要设置 indent）：
    {{
      "content": [
        {{"text": "第一章 概述", "font": "黑体", "font_size": "三号", "bold": true, "align": "center"}},
        {{"text": "这是正文第一段，首行缩进2字符。", "font": "宋体", "font_size": "小四", "indent": 2, "line_spacing": 1.5}},
        {{"text": "这是正文第二段，也必须设置 indent: 2。", "font": "宋体", "font_size": "小四", "indent": 2, "line_spacing": 1.5}},
        {{"text": "这是正文第三段，同样设置 indent: 2。", "font": "宋体", "font_size": "小四", "indent": 2, "line_spacing": 1.5}}
      ],
      "filename": "报告.docx"
    }}

    【修改Word示例】:
    修改全部：{{"file_path": "xxx.docx", "target_style": "all", "font_name": "黑体", "font_size": "三号"}}
    只改标题：{{"file_path": "xxx.docx", "target_style": "title", "font_name": "黑体", "font_size": "三号", "bold": true, "align": "center"}}
    只改正文：{{"file_path": "xxx.docx", "target_style": "body", "font_name": "宋体", "font_size": "四号", "indent": 2, "line_spacing": 1.5}}
    精细修改：{{"file_path": "xxx.docx", "target_style": "body", "font_name": "仿宋", "font_size": "小四", "color": "#333333", "indent": 2, "line_spacing": 1.5}}
    - indent: ⚠️ 首行缩进（字符数）。每个正文段落都必须单独设置此字段！

    【Word字段说明】:
    - font: 字体名称（如 黑体、宋体、仿宋、Arial）
    - font_size: 字号（支持中文：初号、一号、三号、四号、五号等，或数字：12、14）
    - align: 对齐方式（left/center/right/justify）
    - bold: 是否加粗（true/false）
    - italic: 是否斜体（true/false）
    - underline: 是否下划线（true/false）
    - color: 颜色（如 #FF0000 或 red）
    - indent: ⚠️ 首行缩进（字符数）。每个正文段落都必须单独设置此字段！
    - line_spacing: 行间距（如 1.5）

    【注意事项】:
    - 如果文档是表格格式，直接提取表格数据
    - 如果文档是段落文本，根据用户需求提取关键信息
    - 提取的数据必须是 JSON 数组格式
    - 字段名使用中文
    - 如果文档内容中没有用户需要的信息，请如实告知
    - ⚠️ 用户要求首行缩进时，必须在 content 数组中的【每一个】正文段落都设置 "indent": 2
    - ⚠️ 不要用空格代替 indent 参数，必须使用 "indent": 2 字段

    **工具使用指南**:
    - process_excel: 对已有表格数据进行排序/sum/average/count
    - create_excel: 生成Excel文件（支持中文字号、水平/垂直对齐、行高、列宽）
    - create_word: 生成Word文档（支持中文字号、首行缩进、行间距）
    - modify_word: 修改Word样式（支持区分标题/正文，支持字体/字号/颜色/对齐/缩进/行距）
    - modify_excel: 修改Excel样式

    **Excel样式支持**:
    - 字体、字号（支持中文字号或数字）
    - 水平对齐（left/center/right）
    - 垂直对齐（top/center/bottom）
    - 表头背景色（header_color，如 #4F81BD）
    - 表头文字颜色（header_text_color，如 #FFFFFF）
    - 行高（row_height，单位：磅）
    - 列宽（column_widths，单位：字符数）

    请严格按照工具参数要求调用。生成文件后，告知用户文件路径。
    """

            # 安全提取 role 和 content，过滤所有其他字段
            messages_for_agent = []
            for msg in chat_history:
                if isinstance(msg, dict):
                    # 从字典中安全提取
                    role = msg.get('role')
                    content = msg.get('content')
                    if role and content:
                        messages_for_agent.append({
                            'role': str(role),
                            'content': str(content)
                        })
                elif isinstance(msg, list) and len(msg) >= 2:
                    # 兼容旧格式 [user_msg, assistant_msg]
                    if msg[0]:
                        messages_for_agent.append({'role': 'user', 'content': str(msg[0])})
                    if msg[1]:
                        messages_for_agent.append({'role': 'assistant', 'content': str(msg[1])})
            
            # 如果 chat_history 为空，尝试从 session 恢复
            if not messages_for_agent and session.get('messages'):
                for msg in session['messages']:
                    if isinstance(msg, dict):
                        role = msg.get('role')
                        content = msg.get('content')
                        if role and content:
                            messages_for_agent.append({
                                'role': str(role),
                                'content': str(content)
                            })
            
            
            print(f"💬 Agent 消息序列条数: {len(messages_for_agent)}")
            
            # 如果 chat_history 为空，尝试从 session 恢复
            if not messages_for_agent and session.get('messages'):
                for msg in session['messages']:
                    messages_for_agent.append({'role': msg['role'], 'content': msg['content']})

            bot = create_agent_for_task(task_type, system_prompt)

            yield f"🤖 正在处理...", task_label, model_name, f"{len(chat_history)}", ""

            # Agent 流式运行
            last_response = ""
            for chunk in bot.run(messages=messages_for_agent):
                for item in chunk:
                    if isinstance(item, dict):
                        if 'content' in item:
                            current_response = item['content']
                            if len(current_response) > len(last_response):
                                new_content = current_response[len(last_response):]
                                last_response = current_response
                                if new_content and not new_content.startswith('{"status"'):
                                    display = new_content[:500] + "..." if len(new_content) > 500 else new_content
                                    yield display, task_label, model_name, f"{len(chat_history)}", ""
                        elif 'function_call' in item:
                            tool_name = item['function_call'].get('name', '')
                            yield f"🔧 调用工具: {tool_name}...", task_label, model_name, f"{len(chat_history)}", ""

            final_response = last_response
            #  提取生成的文件路径（从 JSON 中解析）
            generated_file = ""
            print(f"📁 开始提取文件路径，final_response 长度: {len(final_response)}")
            print(f"📁 final_response 内容: {final_response[:500]}")
            
            # 1：从 JSON 中解析 file 字段
            try:
                if final_response.startswith('{"status"'):
                    result_json = json.loads(final_response)
                    if 'file' in result_json:
                        generated_file = result_json['file']
                        print(f"📁 从 JSON 解析到文件: {generated_file}")
            except Exception as e:
                print(f"⚠️ JSON 解析失败: {e}")
            
            # 2：如果 JSON 解析失败，尝试正则匹配（支持双反斜杠）
            if not generated_file:
                file_match = re.search(r'([A-Za-z]:\\{1,2}[^\s]+?\.(?:docx|xlsx))', final_response)
                if file_match:
                    generated_file = file_match.group(1).replace('\\\\', '\\')
                    print(f"📁 从正则匹配到文件: {generated_file}")
            
            # 3：如果还是没有，从 output 目录找最新文件
            if not generated_file:
                session_dir = os.path.join(Config.OUTPUT_DIR, session_id)
                if os.path.exists(session_dir):
                    files = os.listdir(session_dir)
                    if files:
                        files.sort(key=lambda f: os.path.getmtime(os.path.join(session_dir, f)), reverse=True)
                        generated_file = os.path.join(session_dir, files[0])
                        print(f"📁 从 output 目录找到最新文件: {generated_file}")
            
            print(f"📁 最终 generated_file: {generated_file}")

            # 同步 chat_history 到 session 持久化存储
            session['messages'] = []
            for msg in chat_history:
                if isinstance(msg, dict) and 'role' in msg and 'content' in msg:
                    session['messages'].append({'role': msg['role'], 'content': msg['content']})
            self.session_manager.update_session(session_id, {'messages': session['messages']})

            # 返回文件信息
            if generated_file and os.path.exists(generated_file):
                print(f"📁 yield 文件信息: {generated_file}")  # 调试打印
                yield "___FILE___", task_label, model_name, f"{len(chat_history)}", generated_file
            else:
                yield "", task_label, model_name, f"{len(chat_history)}", ""

        except Exception as e:
            yield f"❌ 处理失败: {str(e)}", task_label, model_name, f"{len(chat_history)}", ""
    
    def clear_history(self, session_id: str) -> str:
        session = self.session_manager.get_session(session_id)
        if session:
            session['messages'] = []
            self.session_manager.update_session(session_id, {'messages': []})
            return "✅ 对话已清空"
        return "⚠️ 会话不存在"


# ==========================================
# 9. 初始化助手
# ==========================================
assistant = OfficeAssistant()


# ==========================================
# 10. 自定义 CSS（仿 DeepSeek 风格）
# ==========================================
CUSTOM_CSS = """
/* 全局样式 */
body {
    background-color: #f7f8fa;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
}

/* 主容器 */
.gradio-container {
    max-width: 900px !important;
    margin: 0 auto !important;
    padding: 20px !important;
}

/* 顶部标题 */
.app-header {
    text-align: center;
    padding: 20px 0 10px 0;
    border-bottom: 1px solid #e5e7eb;
    margin-bottom: 20px;
}

.app-header h1 {
    font-size: 28px;
    font-weight: 700;
    color: #1f2937;
    margin: 0;
}

.app-header .subtitle {
    color: #6b7280;
    font-size: 14px;
    margin-top: 4px;
}

/* 文件状态栏 */
.file-status-bar {
    background: #ffffff;
    border-radius: 12px;
    padding: 12px 20px;
    margin-bottom: 12px;
    border: 1px solid #e5e7eb;
    display: flex;
    align-items: center;
    justify-content: space-between;
    flex-wrap: wrap;
    font-size: 14px;
    color: #1f2937;
}

/* 状态信息栏 */
.status-bar {
    background: #f1f3f4;
    border-radius: 8px;
    padding: 8px 16px;
    margin-bottom: 16px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    flex-wrap: wrap;
    font-size: 13px;
    color: #4b5563;
}

.status-bar .status-item {
    display: flex;
    align-items: center;
    gap: 4px;
}

.status-bar .status-item .label {
    color: #9ca3af;
}

.status-bar .status-item .value {
    font-weight: 500;
    color: #1f2937;
}

/* 对话区域 */
.chat-container {
    background: #ffffff;
    border-radius: 16px;
    border: 1px solid #e5e7eb;
    overflow: hidden;
    min-height: 400px;
}

/* 输入区域 */
.input-container {
    background: #ffffff;
    border-radius: 16px;
    border: 1px solid #e5e7eb;
    padding: 12px 16px;
    margin-top: 16px;
    display: flex;
    flex-direction: column;
}

.input-row {
    display: flex;
    align-items: flex-end;
    gap: 12px;
}

.input-row textarea {
    flex: 1;
    border: none !important;
    box-shadow: none !important;
    resize: none;
    font-size: 15px;
    padding: 8px 0;
    line-height: 1.6;
    min-height: 52px;
    max-height: 200px;
    background: transparent !important;
}

.input-row textarea:focus {
    outline: none;
    border: none !important;
    box-shadow: none !important;
}

.input-row .upload-btn {
    flex-shrink: 0;
    padding: 8px 12px;
    border-radius: 8px;
    border: 1px solid #e5e7eb;
    background: #f9fafb;
    cursor: pointer;
    font-size: 18px;
    transition: all 0.2s;
}

.input-row .upload-btn:hover {
    background: #f1f3f4;
}

.input-row .send-btn {
    flex-shrink: 0;
    padding: 8px 24px;
    border-radius: 8px;
    border: none;
    background: #2b7ef0;
    color: #ffffff;
    font-weight: 600;
    font-size: 15px;
    cursor: pointer;
    transition: all 0.2s;
}

.input-row .send-btn:hover {
    background: #1a6cd9;
}

.input-row .send-btn:disabled {
    background: #9ca3af;
    cursor: not-allowed;
}

/* 快捷指令 */
.quick-actions {
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
    margin-top: 10px;
    padding-top: 10px;
    border-top: 1px solid #f1f3f4;
}

.quick-actions .quick-btn {
    padding: 4px 14px;
    border-radius: 16px;
    border: 1px solid #e5e7eb;
    background: #f9fafb;
    font-size: 13px;
    color: #4b5563;
    cursor: pointer;
    transition: all 0.2s;
}

.quick-actions .quick-btn:hover {
    background: #e8f0fe;
    border-color: #2b7ef0;
    color: #2b7ef0;
}

/* 底部信息 */
.footer {
    text-align: center;
    padding: 16px 0 8px 0;
    font-size: 13px;
    color: #9ca3af;
    border-top: 1px solid #e5e7eb;
    margin-top: 20px;
}

/* 响应式 */
@media (max-width: 640px) {
    .file-status-bar {
        flex-direction: column;
        gap: 8px;
    }
    .status-bar {
        flex-direction: column;
        gap: 4px;
        align-items: flex-start;
    }
    .input-row {
        flex-wrap: wrap;
    }
    .quick-actions {
        gap: 4px;
    }
}
"""


# ==========================================
#11. Gradio 界面
# ==========================================
def create_interface():
    # 将 theme 和 css 移到 launch() 中
    with gr.Blocks(
        title="智能办公助手",
    ) as demo:

        # 顶部标题
        gr.HTML("""
        <div class="app-header">
            <h1>🤖 智能办公助手</h1>
            <div class="subtitle">支持：文档检索问答 · 生成Excel/Word · 修改样式</div>
        </div>
        """)

        # 状态存储
        session_state = gr.State({"id": None})

        # 文件状态栏
        with gr.Row(equal_height=False):
            file_status = gr.Textbox(
                label="",
                value="📎 未加载文件",
                interactive=False,
                container=False,
                elem_classes="file-status-bar"
            )

        # 状态信息栏
        with gr.Row(equal_height=False):
            with gr.Group(elem_classes="status-bar"):
                task_status = gr.Textbox(
                    label="",
                    value="📌 待识别",
                    interactive=False,
                    container=False,
                    scale=1
                )
                model_status = gr.Textbox(
                    label="",
                    value="🧠 qwen-turbo",
                    interactive=False,
                    container=False,
                    scale=1
                )
                round_status = gr.Textbox(
                    label="",
                    value="💬 第 0 轮",
                    interactive=False,
                    container=False,
                    scale=1
                )

        # 移除 show_copy_button 参数（Gradio 6.0 中已移除）
        chatbot = gr.Chatbot(
            label="",
            height=500,
            elem_classes="chat-container",
        )

        # 输入区域
        with gr.Group(elem_classes="input-container"):
            with gr.Row(elem_classes="input-row", equal_height=False):
                upload_btn = gr.UploadButton(
                    "📎",
                    file_types=[".pdf", ".docx", ".xlsx", ".txt", ".md"],
                    elem_classes="upload-btn",
                    scale=0,
                )
                msg_input = gr.Textbox(
                    label="",
                    placeholder="输入指令，Enter 发送，Shift+Enter 换行...",
                    lines=1,
                    max_lines=6,
                    container=False,
                    scale=9,
                    elem_classes="input-textarea"
                )
                send_btn = gr.Button(
                    "发送",
                    variant="primary",
                    elem_classes="send-btn",
                    scale=1,
                )

            # 快捷指令
            with gr.Row(elem_classes="quick-actions"):
                quick_summary = gr.Button("📄 总结文档", elem_classes="quick-btn", scale=0)
                quick_extract = gr.Button("📊 提取数据生成表格", elem_classes="quick-btn", scale=0)
                quick_format = gr.Button("📝 调整格式", elem_classes="quick-btn", scale=0)
                quick_modify = gr.Button("✏️ 修改样式", elem_classes="quick-btn", scale=0)
                quick_analyze = gr.Button("🔍 分析数据", elem_classes="quick-btn", scale=0)
                quick_clear = gr.Button("🗑️ 清空对话", elem_classes="quick-btn", scale=0)

        # 底部信息
        gr.HTML("""
        <div class="footer">
            支持: PDF · Word · Excel · TXT · MD
        </div>
        """)

        # ========== 事件绑定 ==========

        def init_session():
            session_id = str(uuid.uuid4())[:8]
            return {"id": session_id}

        demo.load(fn=init_session, outputs=[session_state])

        def handle_upload(file, session_state):
            session_id = session_state.get("id")
            if not session_id:
                session_id = str(uuid.uuid4())[:8]
                session_state["id"] = session_id
            result, status, words = assistant.upload_file(file, session_id)
            return result, status, words, session_state

        upload_btn.upload(
            fn=handle_upload,
            inputs=[upload_btn, session_state],
            outputs=[file_status, file_status, file_status, session_state]
        )

        def respond(message, history, session_state):
            session_id = session_state.get("id")
            if not session_id:
                session_id = str(uuid.uuid4())[:8]
                session_state["id"] = session_id

            if not message or not message.strip():
                return history, "", "", session_state, ""

            if history is None:
                history = []
            
            # 添加用户消息到显示层
            history.append({"role": "user", "content": message})
            history.append({"role": "assistant", "content": ""})

            full_response = ""
            last_file_path = ""

            #  传入 chat_history（即 history）给 process_message
            for chunk, task, model, round_num, file_path in assistant.process_message(
                message, history, session_id
            ):
                if chunk == "___FILE___":
                    last_file_path = file_path
                    continue
                elif chunk:
                    full_response += chunk
                    history[-1]["content"] = full_response
                    yield history, task, model, round_num, session_state, ""

            # 如果有文件，确保显示下载链接
            if last_file_path and os.path.exists(last_file_path):
                file_name = os.path.basename(last_file_path)
                output_dir = os.path.join(Config.OUTPUT_DIR, session_id)
                os.makedirs(output_dir, exist_ok=True)
                dest_path = os.path.join(output_dir, file_name)
                shutil.copy2(last_file_path, dest_path)
                # 如果助手还没有回复内容，先添加一个提示
                if not history[-1]["content"]:
                    history[-1]["content"] = "✅ 文件已生成！"
                history[-1]["content"] += f"\n\n📄 [{file_name}](file={dest_path})"
                yield history, task, model, round_num, session_state, ""

       
        send_btn.click(
            fn=respond,
            inputs=[msg_input, chatbot, session_state],
            outputs=[chatbot, task_status, model_status, round_status, session_state, msg_input],  # 增加 msg_input，显示重置对话框为空字符串
        )

        msg_input.submit(
            fn=respond,
            inputs=[msg_input, chatbot, session_state],
            outputs=[chatbot, task_status, model_status, round_status, session_state, msg_input],  #  增加 msg_input 
        )
        # 快捷指令
        quick_commands = {
            quick_summary: "请总结这份文档的主要内容",
            quick_extract: "从文档中提取关键数据，生成Excel表格",
            quick_format: "调整文档格式，使其更加规范美观",
            quick_modify: "修改文档样式，设置合适的字体和排版",
            quick_analyze: "分析文档中的数据和关键信息",
        }

        for btn, cmd in quick_commands.items():
            btn.click(
                fn=lambda c=cmd: c,
                inputs=[],
                outputs=[msg_input]
            ).then(
                fn=lambda: None,
                inputs=[],
                outputs=[],
                js="() => { document.querySelector('.send-btn').click(); }"
            )

        def clear_chat(history, session_state):
            session_id = session_state.get("id")
            if session_id:
                assistant.clear_history(session_id)
            return [], "📌 待识别", "🧠 qwen-turbo", "💬 第 0 轮", session_state

        quick_clear.click(
            fn=clear_chat,
            inputs=[chatbot, session_state],
            outputs=[chatbot, task_status, model_status, round_status, session_state]
        )

        return demo


# ==========================================
# 12. 主程序入口
# ==========================================
if __name__ == "__main__":
    demo = create_interface()

    parser = argparse.ArgumentParser(description="智能办公助手")
    parser.add_argument("--share", action="store_true", help="开启公网分享")
    parser.add_argument("--port", type=int, default=7860, help="端口号")
    parser.add_argument("--host", type=str, default="127.0.0.1", help="主机地址")
    args = parser.parse_args()

    print("=" * 60)
    print("🤖 智能办公助手 v1.0")
    print(f"  本地访问: http://{args.host}:{args.port}")
    if args.share:
        print("  公网分享: 启动后生成临时链接")
    print("=" * 60)

    demo.queue(default_concurrency_limit=10).launch(
        server_name=args.host,
        server_port=args.port,
        share=args.share,
        debug=False,
        show_error=True,
        allowed_paths=[Config.OUTPUT_DIR],  # 允许访问 outputs 目录
        theme=gr.themes.Default(
            primary_hue="blue",
            neutral_hue="gray",
            font=gr.themes.GoogleFont("Inter")
        ),
        css=CUSTOM_CSS,
    )